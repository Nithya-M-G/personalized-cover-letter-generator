import streamlit as st
import uuid
import ollama
import numpy as np
import easyocr
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image, ImageOps, ImageEnhance
import io

# 1. Page Configuration
st.set_page_config(page_title="Personalized Cover Letter Generator", layout="centered")

# 2. Initialize EasyOCR Reader
@st.cache_resource
def load_ocr_reader():
    return easyocr.Reader(['en'], gpu=False)

reader = load_ocr_reader()

# 3. Extraction Helper (Handles Resume and JD)
def extract_text(uploaded_file):
    extension = uploaded_file.name.split('.')[-1].lower()
    
    if extension in ['jpg', 'jpeg', 'png']:
        img = Image.open(uploaded_file).convert('RGB')
        img = ImageOps.grayscale(img)
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img_array = np.array(img)
        results = reader.readtext(img_array, detail=0)
        return " ".join(results)

    elif extension == 'pdf':
        pdf_reader = PdfReader(uploaded_file)
        return " ".join([p.extract_text() for p in pdf_reader.pages if p.extract_text()])

    elif extension == 'docx':
        doc = Document(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs])
    
    else:
        return uploaded_file.read().decode("utf-8")

# 4. Session State Setup
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = {}
if "current_chat_id" not in st.session_state:
    st.session_state.current_chat_id = None

# 5. Sidebar Navigation
with st.sidebar:
    st.header("Your Chats")
    if st.button("➕ New Chat"):
        new_id = str(uuid.uuid4())
        st.session_state.chat_sessions[new_id] = {
            "messages": [], 
            "resume_text": None,
            "resume_filename": None,
            "jd_text": None,
            "jd_filename": None
        }
        st.session_state.current_chat_id = new_id
        st.rerun()

    st.divider()
    for chat_id, session_data in st.session_state.chat_sessions.items():
        msg = session_data["messages"]
        label = msg[0]["content"][:20] + "..." if msg else f"Chat {chat_id[:4]}"
        if st.button(label, key=chat_id):
            st.session_state.current_chat_id = chat_id
            st.rerun()

# 6. Main UI
st.title("Personalized Cover Letter Generator")

if st.session_state.current_chat_id is None:
    st.info("👋 Welcome! Start a new chat to upload your documents.")
    st.stop()

current_session = st.session_state.chat_sessions[st.session_state.current_chat_id]

# --- Dual Upload Section ---
col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Resume")
    if current_session["resume_text"] is None:
        res_file = st.file_uploader("Upload Resume", type=['pdf', 'docx', 'txt', 'jpg', 'png', 'jpeg'], key="res")
        if res_file:
            with st.spinner("Extracting Resume..."):
                current_session["resume_text"] = extract_text(res_file)
                current_session["resume_filename"] = res_file.name
                st.rerun()
    else:
        st.success(f"📄 {current_session['resume_filename']}")
        if st.button("Reset Resume"):
            current_session["resume_text"] = None
            st.rerun()

with col2:
    st.subheader("2. Job Description")
    if current_session["jd_text"] is None:
        jd_file = st.file_uploader("Upload JD", type=['pdf', 'docx', 'txt', 'jpg', 'png', 'jpeg'], key="jd")
        if jd_file:
            with st.spinner("Extracting Job Description..."):
                current_session["jd_text"] = extract_text(jd_file)
                current_session["jd_filename"] = jd_file.name
                st.rerun()
    else:
        st.success(f"💼 {current_session['jd_filename']}")
        if st.button("Reset JD"):
            current_session["jd_text"] = None
            st.rerun()

st.divider()

# 7. Chat Interface
for message in current_session["messages"]:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# 8. Llama 3 Chat Logic (Grounded on Resume + JD)
if prompt := st.chat_input("Ask to generate or refine your cover letter..."):
    with st.chat_message("user"):
        st.markdown(prompt)
    current_session["messages"].append({"role": "user", "content": prompt})

    with st.chat_message("assistant"):
        if current_session["resume_text"] and current_session["jd_text"]:
            res_placeholder = st.empty()
            full_res = ""
            try:
                # 
                system_instr = (
                    "You are a Career Coach. Write a formal cover letter by matching the Resume to the JD. "
                    "STRICT RULE: Only use facts from the Resume. Use the JD to highlight matching skills. "
                    "DO NOT invent experience. DO NOT use brackets [ ]."
                )
                user_content = (
                    f"RESUME DATA:\n{current_session['resume_text']}\n\n"
                    f"JOB DESCRIPTION DATA:\n{current_session['jd_text']}\n\n"
                    f"USER REQUEST: {prompt}"
                )

                # Use 'phi3' if 'llama3' hits memory errors (4.6GiB vs 3.1GiB)
                stream = ollama.chat(
                    model='llama3', 
                    messages=[{'role': 'system', 'content': system_instr},
                              {'role': 'user', 'content': user_content}],
                    stream=True
                )
                for chunk in stream:
                    full_res += chunk['message']['content']
                    res_placeholder.markdown(full_res + "▌")
                res_placeholder.markdown(full_res)
                current_session["messages"].append({"role": "assistant", "content": full_res})
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.warning("Please upload BOTH the Resume and Job Description first.")
# 9. Download Functionality
if current_session["messages"]:
    # Find the last assistant message (the cover letter)
    last_assistant_msg = next((msg["content"] for msg in reversed(current_session["messages"]) 
                               if msg["role"] == "assistant"), None)
    
    if last_assistant_msg:
        st.divider()
        st.subheader("📥 Export Your Cover Letter")
        
        # Option 1: Download as .txt
        st.download_button(
            label="Download as Text File",
            data=last_assistant_msg,
            file_name="Cover_Letter.txt",
            mime="text/plain"
        )

        # Option 2: Download as .docx (Professional Format)
        doc = Document()
        doc.add_heading('Cover Letter', 0)
        doc.add_paragraph(last_assistant_msg)
        
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="Download as Word Document",
            data=bio.getvalue(),
            file_name="Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )